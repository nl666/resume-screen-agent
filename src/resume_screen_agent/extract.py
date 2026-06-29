from __future__ import annotations

from pathlib import Path
from dataclasses import dataclass


SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".jsonl", ".docx", ".pdf"}


@dataclass(frozen=True)
class ExtractedText:
    text: str
    char_count: int
    status: str


def read_text(path: str | Path) -> str:
    return read_text_with_diagnostics(path).text


def read_text_with_diagnostics(path: str | Path) -> ExtractedText:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix in {".txt", ".md", ".csv", ".jsonl"}:
        text = file_path.read_text(encoding="utf-8")
    elif suffix == ".docx":
        text = _read_docx(file_path)
    elif suffix == ".pdf":
        text = _read_pdf(file_path)
    else:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type: {suffix}. Supported: {supported}")

    char_count = len(text.strip())
    if char_count == 0:
        status = "empty"
    elif suffix == ".pdf" and char_count < 300:
        status = "too_short_check_ocr"
    else:
        status = "ok"
    return ExtractedText(text=text, char_count=char_count, status=status)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Reading .docx requires python-docx. Run: pip install -r requirements.txt") from exc

    document = Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Reading .pdf requires pypdf. Run: pip install -r requirements.txt") from exc

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()
